#!/usr/bin/env python3
"""Render the Markdown handbooks in docs/ to .docx for people who read Word.

The .docx files are a convenience copy — the Markdown is the source of truth.
Re-run this after editing any of them:

    python3 docs/build_docx.py            # rebuild all four
    python3 docs/build_docx.py FORMS.md   # rebuild one

Needs python-docx (`pip3 install python-docx`) and nothing else, so it works
without pandoc. The styling matches the original hand-built exports: Heading
1-3, Intense Quote for blockquotes, Light Grid Accent 1 for tables, Consolas
for code, and centred screenshots with an italic caption underneath.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCS = Path(__file__).resolve().parent
DEFAULT_SOURCES = ["ARCHITECTURE.md", "AUTHENTICATION.md", "MATTERS.md", "FORMS.md"]

CODE_FONT = "Consolas"
CODE_SIZE = Pt(9.5)          # sz 19, matches the existing exports
BLOCK_CODE_SIZE = Pt(9)      # sz 18
LINK_COLOR = RGBColor(0x05, 0x63, 0xC1)
IMAGE_WIDTH = Inches(6.2)

# `code`, **bold**, *italic*, [text](href) — in that precedence order. Bold is
# matched non-greedily so a bold span may itself contain an italic or a code
# span, which several of these docs rely on.
INLINE_RE = re.compile(
    r"(?P<code>`[^`]+`)"
    r"|(?P<bold>\*\*.+?\*\*)"
    r"|(?P<link>\[[^\]]+\]\([^)]+\))"
    r"|(?P<italic>\*[^*\n]+\*)"
)


# --------------------------------------------------------------------------
# inline runs
# --------------------------------------------------------------------------

def parse_inline(text):
    """Split a line into (text, bold, italic, code, href) tuples."""
    out, pos = [], 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            out.append((text[pos:m.start()], False, False, False, None))
        if m.group("code"):
            out.append((m.group()[1:-1], False, False, True, None))
        elif m.group("bold"):
            # Bold can wrap code spans: **`id`, `uid` and `sid`**
            for t, _, i, c, h in parse_inline(m.group()[2:-2]):
                out.append((t, True, i, c, h))
        elif m.group("italic"):
            out.append((m.group()[1:-1], False, True, False, None))
        else:
            label, href = re.match(r"\[([^\]]+)\]\(([^)]+)\)", m.group()).groups()
            out.append((label, False, False, False, href))
        pos = m.end()
    if pos < len(text):
        out.append((text[pos:], False, False, False, None))
    return [r for r in out if r[0]]


def strip_markup(text):
    """Headings and captions are written as plain text, without the markers."""
    return "".join(t for t, *_ in parse_inline(text))


def add_hyperlink(paragraph, label, href):
    part = paragraph.part
    r_id = part.relate_to(
        href,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(color)
    props.append(underline)
    run.append(props)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    link.append(run)
    paragraph._p.append(link)


def write_runs(paragraph, text):
    for chunk, bold, italic, code, href in parse_inline(text):
        if href:
            add_hyperlink(paragraph, chunk, href)
            continue
        run = paragraph.add_run(chunk)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = CODE_FONT
            run.font.size = CODE_SIZE


# --------------------------------------------------------------------------
# blocks
# --------------------------------------------------------------------------

def add_code_block(doc, lines, language):
    if language == "mermaid":
        label = doc.add_paragraph()
        run = label.add_run("Mermaid diagram source:")
        run.italic = True
        run.font.size = BLOCK_CODE_SIZE

    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.2)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), "F2F2F2")
    para._p.get_or_add_pPr().append(shading)

    run = para.add_run()
    run.font.name = CODE_FONT
    run.font.size = BLOCK_CODE_SIZE
    for i, line in enumerate(lines):
        if i:
            run.add_break()
        run.add_text(line)


def add_table(doc, rows):
    header, body = rows[0], rows[2:]      # rows[1] is the |---|---| separator
    table = doc.add_table(rows=0, cols=len(header))
    table.style = "Light Grid Accent 1"
    for cells, is_header in [(header, True)] + [(r, False) for r in body]:
        row = table.add_row().cells
        for cell, text in zip(row, cells):
            para = cell.paragraphs[0]
            if is_header:
                run = para.add_run(strip_markup(text))
                run.bold = True
            else:
                write_runs(para, text)


def add_image(doc, alt, src):
    path = DOCS / src
    if not path.exists():
        doc.add_paragraph(f"[missing image: {src}]")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(path), width=IMAGE_WIDTH)

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(strip_markup(alt))
    run.italic = True
    run.font.size = BLOCK_CODE_SIZE


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def convert(md_path, docx_path):
    lines = md_path.read_text().split("\n")
    doc = Document()
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # fenced code / mermaid
        if stripped.startswith("```"):
            language = stripped[3:].strip()
            i += 1
            block = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            add_code_block(doc, block, language)
            i += 1
            continue

        # horizontal rule -> blank spacer paragraph
        if re.fullmatch(r"-{3,}", stripped):
            doc.add_paragraph()
            i += 1
            continue

        # headings
        heading = re.match(r"(#{1,3})\s+(.*)", stripped)
        if heading:
            level = len(heading.group(1))
            doc.add_heading(strip_markup(heading.group(2)), level=level)
            i += 1
            continue

        # standalone image
        image = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image:
            add_image(doc, image.group(1), image.group(2))
            i += 1
            continue

        # blockquote: each run of lines becomes its own Intense Quote paragraph,
        # and any table rows inside it stay as literal pipe text.
        if stripped.startswith(">"):
            while i < len(lines) and lines[i].strip().startswith(">"):
                chunk = []
                while i < len(lines) and lines[i].strip().startswith(">"):
                    body = lines[i].strip()[1:].strip()
                    if not body:
                        i += 1
                        break
                    if body.startswith("|") and chunk:
                        break
                    chunk.append(body)
                    i += 1
                    if body.startswith("|"):
                        break
                if chunk:
                    para = doc.add_paragraph(style="Intense Quote")
                    write_runs(para, " ".join(chunk))
            continue

        # table
        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            if len(rows) >= 2:
                add_table(doc, rows)
            continue

        # list item (one level of nesting, bulleted or numbered)
        listitem = re.match(r"(\s*)([-*]|\d+\.)\s+(.*)", line)
        if listitem:
            indent, marker, text = listitem.groups()
            nested = len(indent) >= 2
            bulleted = marker in ("-", "*")
            style = "List Bullet" if bulleted else "List Number"
            if nested:
                style += " 2"
            i += 1
            # wrapped continuation lines belong to the same item
            while i < len(lines):
                nxt = lines[i]
                if not nxt.strip() or re.match(r"\s*([-*]|\d+\.)\s+", nxt):
                    break
                if not nxt.startswith(" "):
                    break
                text += " " + nxt.strip()
                i += 1
            write_runs(doc.add_paragraph(style=style), text)
            continue

        # plain paragraph: join wrapped lines
        chunk = []
        while i < len(lines):
            nxt = lines[i]
            if not nxt.strip():
                break
            if re.match(r"\s*([-*]|\d+\.)\s+|#{1,3}\s|>|\||```", nxt.strip()):
                break
            if re.fullmatch(r"-{3,}", nxt.strip()):
                break
            chunk.append(nxt.strip())
            i += 1
        if chunk:
            write_runs(doc.add_paragraph(), " ".join(chunk))

    doc.save(docx_path)
    return docx_path


def main():
    sources = sys.argv[1:] or DEFAULT_SOURCES
    for name in sources:
        md = DOCS / Path(name).name
        if not md.exists():
            sys.exit(f"no such file: {md}")
        out = convert(md, md.with_suffix(".docx"))
        print(f"{md.name} -> {out.name}")


if __name__ == "__main__":
    main()
