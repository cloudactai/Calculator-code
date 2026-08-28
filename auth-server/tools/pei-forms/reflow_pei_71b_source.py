#!/usr/bin/env python3
"""Rebuild PEI Form 71B as a genuinely reflowed Word source.

The published Form 71B is a legacy ``.doc``.  LibreOffice renders that source
faithfully, including its cramped opening line, no writing room after item
(c), an ambiguous signing block, and a duplicated final ``Registrar`` label.
Those defects cannot be fixed safely by masking and drawing over the rendered
PDF: replacement text uses a different PDF font and can overlap the original
text at browser zoom levels.

This tool validates the legal copy against the converted government source,
then lays the same copy out in a clean DOCX using one native Times New Roman
typographic system.  LibreOffice subsequently renders that DOCX, so every
line, label, and paragraph participates in Word's normal page flow.

    python3 reflow_pei_71b_source.py \
      --source ../../form-template-export/_incoming_pei/PEISC_71B_source.doc \
      --output ../../form-template-export/_incoming_pei/PEISC_71B_reflow.docx

Requires python-docx, which the repository already uses in docs/build_docx.py.
"""

import argparse
import os
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SOFFICE = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
FONT = "Times New Roman"
BODY_SIZE = Pt(10)
CAPTION_SIZE = Pt(7)

OPENING = (
    "I, (full name), of (address), acknowledge that I am indebted to Her "
    "Majesty the Queen in the amount of $_____, to be collected from me in "
    "any manner that an order for the payment of money may be enforced, if I "
    "fail (where the person enters into the recognizance as surety for another "
    "person, substitute if (name) fails) to abide by any of the following "
    "conditions:"
)
CONDITION = (
    "1. That until (date) (or until the court orders otherwise) I shall not "
    "(where the person enters into the recognizance as a surety, substitute "
    "(name) shall not), "
)
CERTIFICATION = (
    "This recognizance was entered into before me in accordance with the "
    "order of the court made on (date)."
)


def set_run_font(run, size=BODY_SIZE, italic=False, bold=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = size
    run.italic = italic
    run.bold = bold
    return run


def set_paragraph(paragraph, *, align=None, before=0, after=0,
                  exact=12, left=0, first=None, keep=False):
    fmt = paragraph.paragraph_format
    if align is not None:
        paragraph.alignment = align
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.left_indent = Pt(left)
    if first is not None:
        fmt.first_line_indent = Pt(first)
    if exact:
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(exact)
    fmt.keep_with_next = keep
    return paragraph


def clear_table_borders(table):
    props = table._tbl.tblPr
    borders = props.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        props.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start),
                        ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_bottom_border(cell, size=8):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")


def set_cell_width(cell, points):
    cell.width = Pt(points)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    tc_w.set(qn("w:w"), str(round(points * 20)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_indent(table, points):
    props = table._tbl.tblPr
    indent = props.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        props.append(indent)
    indent.set(qn("w:w"), str(round(points * 20)))
    indent.set(qn("w:type"), "dxa")


def set_table_fixed(table):
    table.autofit = False
    layout = table._tbl.tblPr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table._tbl.tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_table_widths(table, widths):
    """Set both cell widths and tblGrid widths; LibreOffice reads the grid."""
    grid = table._tbl.tblGrid
    for column, width in zip(grid.gridCol_lst, widths):
        column.set(qn("w:w"), str(round(width * 20)))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def remove_cell_paragraph_padding(cell):
    for paragraph in cell.paragraphs:
        set_paragraph(paragraph, exact=11)


def add_text(paragraph, text, *, size=BODY_SIZE, italic=False, bold=False):
    return set_run_font(paragraph.add_run(text), size, italic, bold)


def add_underlined_blank(paragraph, points):
    # A native underlined run participates in Word's line flow; it is not a
    # PDF line or replacement text drawn after the page has been rendered.
    run = add_text(paragraph, "\u00a0" * max(2, round(points / 2.5)))
    run.underline = True
    return run


def add_segmented_paragraph(doc, segments, **kwargs):
    paragraph = set_paragraph(doc.add_paragraph(), **kwargs)
    for text, italic in segments:
        if text.startswith("@@BLANK:") and text.endswith("@@"):
            add_underlined_blank(paragraph, float(text[8:-2]))
        else:
            add_text(paragraph, text, italic=italic)
    return paragraph


def add_caption(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = set_paragraph(cell.paragraphs[0], align=align, exact=8)
    add_text(paragraph, text, size=CAPTION_SIZE, italic=True)


def add_opening_fields(doc):
    # Fixed columns keep the address line completely clear of "of" at every
    # zoom.  The 172pt address rule is intentionally longer than the 94pt name
    # rule, while leaving enough room for punctuation on the same source line.
    widths = (18, 94, 28, 172, 10)
    table = doc.add_table(rows=2, cols=len(widths))
    clear_table_borders(table)
    set_table_fixed(table)
    set_table_widths(table, widths)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(14 if row is table.rows[0] else 8)
        for cell, width in zip(row.cells, widths):
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            remove_cell_paragraph_padding(cell)

    values = ("I,", "", ", of", "", ",")
    for cell, value in zip(table.rows[0].cells, values):
        paragraph = set_paragraph(cell.paragraphs[0], exact=12)
        add_text(paragraph, value)
    set_cell_bottom_border(table.rows[0].cells[1])
    set_cell_bottom_border(table.rows[0].cells[3])
    add_caption(table.rows[1].cells[1], "(full name)")
    add_caption(table.rows[1].cells[3], "(address)")
    return table


def add_answer_line(doc, marker, lead, caption, suffix=""):
    table = doc.add_table(rows=2, cols=4)
    clear_table_borders(table)
    set_table_fixed(table)
    set_table_indent(table, 23)
    # The lead and rule share the available width without any PDF overlay.
    lead_width = 160 if marker == "(a)" else 270
    rule_width = 220 if marker == "(a)" else 110
    widths = (28, lead_width, rule_width, 10)
    set_table_widths(table, widths)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(13 if row is table.rows[0] else 8)
        for cell, width in zip(row.cells, widths):
            set_cell_margins(cell)
            remove_cell_paragraph_padding(cell)
    for cell, value, italic in zip(
        table.rows[0].cells,
        (marker, lead, "", suffix),
        (False, False, False, False),
    ):
        p = set_paragraph(cell.paragraphs[0], exact=11)
        add_text(p, value, italic=italic)
    set_cell_bottom_border(table.rows[0].cells[2])
    add_caption(table.rows[1].cells[2], caption)
    return table


def add_terms_area(doc):
    p = set_paragraph(doc.add_paragraph(), left=51, exact=11, before=1)
    add_text(p, "(c)")
    add_text(p, "   (include any other terms).", italic=True)

    table = doc.add_table(rows=1, cols=1)
    clear_table_borders(table)
    set_table_fixed(table)
    set_table_indent(table, 51)
    set_table_widths(table, (366,))
    set_cell_margins(table.cell(0, 0))
    row = table.rows[0]
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    row.height = Pt(62)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    set_paragraph(cell.paragraphs[0], exact=8)
    # This is the rule that visually belongs to the text area.  The overlay
    # mapping occupies the open band above it.
    set_cell_bottom_border(cell)
    return table


def add_signing_pair(doc, *, registrar=False, before=3):
    table = doc.add_table(rows=3 if registrar else 2, cols=3)
    clear_table_borders(table)
    set_table_fixed(table)
    set_table_indent(table, 51)
    widths = (153, 60, 153)
    set_table_widths(table, widths)
    for index, row in enumerate(table.rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(12 if index == 0 else 9)
        for cell, width in zip(row.cells, widths):
            set_cell_margins(cell)
            remove_cell_paragraph_padding(cell)
    for cell in table.rows[0].cells:
        cell.paragraphs[0].paragraph_format.space_before = Pt(before)
    set_cell_bottom_border(table.rows[0].cells[0])
    set_cell_bottom_border(table.rows[0].cells[2])
    add_caption(table.rows[1].cells[0], "(Date)")
    add_caption(table.rows[1].cells[2], "(Signature)", WD_ALIGN_PARAGRAPH.RIGHT)
    if registrar:
        p = set_paragraph(table.rows[2].cells[2].paragraphs[0],
                          align=WD_ALIGN_PARAGRAPH.RIGHT, exact=10)
        add_text(p, "Registrar")
    return table


def add_order_date(doc):
    table = doc.add_table(rows=2, cols=2)
    clear_table_borders(table)
    set_table_fixed(table)
    set_table_widths(table, (153, 263))
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(10)
        for cell in row.cells:
            set_cell_margins(cell)
            remove_cell_paragraph_padding(cell)
    set_cell_bottom_border(table.rows[0].cells[0])
    add_text(table.rows[0].cells[1].paragraphs[0], ".")
    add_caption(table.rows[1].cells[0], "(Date of court order)")
    return table


def convert_to_docx(source, out_dir):
    source = Path(source)
    if source.suffix.lower() == ".docx":
        return source
    subprocess.run(
        [SOFFICE, "--headless", "--convert-to", "docx", "--outdir",
         str(out_dir), str(source)],
        check=True,
        capture_output=True,
    )
    converted = Path(out_dir) / (source.stem + ".docx")
    if not converted.exists():
        raise RuntimeError("LibreOffice did not produce %s" % converted)
    return converted


def source_text(source):
    with tempfile.TemporaryDirectory(prefix="pei71b-source-") as temp:
        converted = convert_to_docx(source, temp)
        doc = Document(str(converted))
        return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())


def validate_source(source):
    text = re.sub(r"\s+", " ", source_text(source)).strip()
    required = (
        "FORM 71 B", "RECOGNIZANCE", OPENING, CONDITION,
        "(a) enter on the premises at (address);",
        "(b) speak to, telephone, write to or otherwise communicate with (name);",
        "(c) (include any other terms).", CERTIFICATION,
    )
    missing = [
        phrase for phrase in required
        if re.sub(r"\s+", " ", phrase).strip() not in text
    ]
    if missing:
        raise ValueError("Source copy changed; refusing to reflow: %r" % missing)


def build(source, output):
    validate_source(source)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(1.15)
    section.right_margin = Inches(1.15)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.28)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = BODY_SIZE
    normal.paragraph_format.space_after = Pt(0)

    first = doc.add_paragraph()
    set_paragraph(first, align=WD_ALIGN_PARAGRAPH.CENTER, exact=13,
                  after=0, keep=True)
    add_text(first, "FORM 71 B", size=Pt(12), bold=True)
    p = set_paragraph(doc.add_paragraph(), align=WD_ALIGN_PARAGRAPH.CENTER,
                      exact=13, after=12, keep=True)
    add_text(p, "RECOGNIZANCE", size=Pt(12), bold=True)

    p = set_paragraph(doc.add_paragraph(), align=WD_ALIGN_PARAGRAPH.CENTER,
                      exact=12, keep=True)
    add_text(p, "(General heading)", italic=True)
    p = set_paragraph(doc.add_paragraph(), align=WD_ALIGN_PARAGRAPH.CENTER,
                      exact=12, after=5, keep=True)
    add_text(p, "RECOGNIZANCE")

    add_opening_fields(doc)

    add_segmented_paragraph(
        doc,
        (
            ("acknowledge that I am indebted to Her Majesty the Queen in the amount of $", False),
            ("@@BLANK:50@@", False),
            (", to be collected from me in any manner that an order for the payment of money may be enforced, if I fail ", False),
            ("(where the person enters into the recognizance as surety for another person, substitute if (name) fails)", True),
            (" to abide by any of the following conditions:", False),
        ),
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        exact=12,
        after=2,
    )

    add_segmented_paragraph(
        doc,
        (
            ("1.        That until ", False),
            ("@@BLANK:62@@", False),
            (" ", False),
            ("(date)", True),
            (" ", False),
            ("(or until the court orders otherwise)", True),
            (" I shall not ", False),
            ("(where the person enters into the recognizance as a surety, substitute (name) shall not),", True),
        ),
        exact=12,
        left=0,
        first=0,
        after=1,
    )

    p = set_paragraph(doc.add_paragraph(), left=23, exact=11, after=1)
    add_text(p, "(Insert where applicable:)", italic=True)
    add_answer_line(doc, "(a)", "enter on the premises at", "(address)", ";")
    add_answer_line(doc, "(b)", "speak to, telephone, write to or otherwise communicate with", "(name)", ";")
    add_terms_area(doc)
    add_signing_pair(doc, before=2)

    add_segmented_paragraph(
        doc,
        (
            ("This recognizance was entered into before me in accordance with the order of the court made on", False),
        ),
        exact=12,
        before=5,
        after=1,
    )
    add_order_date(doc)
    add_signing_pair(doc, registrar=True, before=2)

    footer = section.footer.paragraphs[0]
    set_paragraph(footer, align=WD_ALIGN_PARAGRAPH.CENTER, exact=10)
    add_text(footer, "1", size=Pt(9))

    output = Path(output)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    build(args.source, args.output)
    print("wrote %s" % os.path.abspath(args.output))


if __name__ == "__main__":
    main()
